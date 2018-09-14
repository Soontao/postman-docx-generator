from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from postman_schema import PostmanObject, Items
from string import Template
import datetime

POSTMAN_COLLECTION_SCHEMA_V2_1 = "https://schema.getpostman.com/json/collection/v2.1.0/collection.json"


class DocumentGenerator(object):

  def __init__(self, metadata: PostmanObject):
    super().__init__()
    self.document = Document()
    self.metadata = metadata
    self._apply_style()
    self._check_data()
    self._document_title()
    self._add_requests()

  def _check_data(self):
    assert len(self.metadata.item) > 0, "collection must have at least one request"
    assert self.metadata.info.schema == POSTMAN_COLLECTION_SCHEMA_V2_1, "only support postman collection schema v2.1"

  def _apply_style(self):
    heading_1 = self.document.styles['Heading 1']
    heading_1.font.size = Pt(25)

  def _document_title(self):
    self.document.add_heading(self.metadata.info.name, 0)
    self.document.add_paragraph(self.metadata.info.description)
    self.document.add_heading("Version", 1)
    self._add_table(rows=[["Version", "Date", "Comments"],
                          [1, datetime.datetime.now().strftime("%Y-%m-%d %H:%M"), "init document"]])

  def _add_requests(self):
    for request in self.metadata.item:
      self._add_api_chapter(request)

  def _add_api_chapter(self, api_item: Items):
    self.document.add_heading(api_item.name, 1)
    self.document.add_paragraph(api_item.request.description)

    # request basic info
    self.document.add_heading("Request", 2)
    self._add_table([["Method", "URL"], [api_item.request.method, api_item.request.url.raw]])

    # add example response
    for example in api_item.response:
      self._add_api_example(example)

  def _add_api_example(self, example):
    pass

  def _add_table(self, rows):
    assert len(rows) > 0, "table must have at list one records"
    column_number = len(rows[0])
    rows_number = len(rows)
    table = self.document.add_table(rows=0, cols=column_number)
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index in range(rows_number):
      row = table.add_row().cells
      for col_index, col_content in enumerate(rows[index]):
        cell = row[col_index]
        cell.text = str(col_content)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

  def save(self, *args, **kwargs):
    self.document.save(*args, **kwargs)
